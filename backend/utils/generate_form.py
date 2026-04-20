#!/usr/bin/env python3
"""
MIT Hostel Visit Form Generator
Generates pixel-perfect DOCX matching original MIT templates.
Usage: python3 generate_form.py <form_type> <json_data> <output_path>
"""

import sys
import json
from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ── Helpers ──────────────────────────────────────────────────────────────────

def border(cell, sides=('top','left','bottom','right'), sz=4, color='000000', val='single'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in sides:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), val)
        tag.set(qn('w:sz'), str(sz))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def no_border(cell):
    border(cell, val='nil')

def cell_margin(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    mar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)

def set_col_width(cell, width_emu):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_emu / 635)))  # EMU to twentieths
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def row_height(row, twips, rule='atLeast'):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(twips))
    h.set(qn('w:hRule'), rule)
    trPr.append(h)

def para_spacing(p, before=0, after=0, line=None):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    if line:
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)

def add_run(para, text, bold=False, size_pt=None, font='Times New Roman'):
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    if size_pt:
        run.font.size = Pt(size_pt)
    return run

def fmt_date(iso):
    if not iso: return ''
    try:
        d = datetime.fromisoformat(iso.replace('Z',''))
        return d.strftime('%d/%m/%Y')
    except: return iso[:10]

def fmt_day(iso):
    if not iso: return ''
    try:
        d = datetime.fromisoformat(iso.replace('Z',''))
        return d.strftime('%A')
    except: return ''

def fmt_time(iso):
    if not iso: return ''
    try:
        d = datetime.fromisoformat(iso.replace('Z',''))
        return d.strftime('%I:%M %p')
    except: return ''

def hostel_to_loc_key(name):
    if not name: return None
    n = name.upper()
    if 'SRTH' in n: return 'srth'
    if 'SVH'  in n: return 'svh'
    if 'TARA' in n: return 'tara'
    if 'SJB'  in n or 'SBP' in n: return 'sjb'
    return None

# ── Page Setup ───────────────────────────────────────────────────────────────

def setup_page(doc):
    sec = doc.sections[0]
    sec.page_width   = Emu(7559055)   # A4 8.27"
    sec.page_height  = Emu(10692630)  # A4 11.69"
    sec.left_margin  = Emu(571500)    # 0.625" (original)
    sec.right_margin = Emu(571500)
    sec.top_margin   = Emu(457200)    # 0.5"
    sec.bottom_margin= Emu(457200)
    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)
    return doc

# ── Header (shared) ─────────────────────────────────────────────────────────

def add_header(doc, committee_line, title_line, hostel_type='Boys'):
    for text, bold, size in [
        ("G.S.Mandal's",                     True,  13),
        ("Maharashtra Institute of Technology,", True, 13),
        ("Chhatrapati Sambhajinagar",         True,  12),
        ("(An Autonomous Institute)",         False, 11),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, 0, 0)
        add_run(p, text, bold=bold, size_pt=size)

    # Horizontal rule
    p = doc.add_paragraph()
    para_spacing(p, 20, 20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.append(pBdr)

    # Committee
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 0, 0)
    add_run(p, committee_line, bold=True, size_pt=12)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 0, 40)
    add_run(p, title_line, bold=False, size_pt=11)

    # Second HR
    p = doc.add_paragraph()
    para_spacing(p, 0, 60)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '000000')
    pBdr.append(top)
    pPr.append(pBdr)

# ── Date/Name row tables ─────────────────────────────────────────────────────

def add_date_name_rows(doc, date_str, day_str, time_str, name_str, dept_str):
    # Widths from original (in EMU): 1714500, 1905000, 892810, 1905000
    # Table 0: Date row
    t0 = doc.add_table(rows=1, cols=4)
    t0.style = 'Table Grid'
    row = t0.rows[0]
    row.cells[0].text = 'Date and Day of visit:'
    row.cells[1].text = f'{date_str}, {day_str}' if date_str else ''
    row.cells[2].text = 'Time slot:-'
    row.cells[3].text = time_str
    widths = [1714500, 1905000, 892810, 1905000]
    for i, cell in enumerate(row.cells):
        set_col_width(cell, widths[i])
        cell_margin(cell, 40, 40, 80, 80)
        for p in cell.paragraphs:
            para_spacing(p, 0, 0)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    doc.add_paragraph()._p.getparent()  # spacer
    # Actually use spacing on a real paragraph
    sp = doc.paragraphs[-1]
    para_spacing(sp, 0, 60)

    # Table 1: Name row  Widths: 1524000, 1905000, 1019810, 1968500
    t1 = doc.add_table(rows=1, cols=4)
    t1.style = 'Table Grid'
    row1 = t1.rows[0]
    row1.cells[0].text = 'Name of Member/s:'
    row1.cells[1].text = name_str
    row1.cells[2].text = 'Department:-'
    row1.cells[3].text = dept_str
    widths1 = [1524000, 1905000, 1019810, 1968500]
    for i, cell in enumerate(row1.cells):
        set_col_width(cell, widths1[i])
        cell_margin(cell, 40, 40, 80, 80)
        for p in cell.paragraphs:
            para_spacing(p, 0, 0)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    sp2 = doc.add_paragraph()
    para_spacing(sp2, 60, 0)

# ── Locations table (shared) ─────────────────────────────────────────────────

LOCS = [('SRTH','srth'), ('SVH','svh'), ('TARA','tara'), ('SJB/SBP','sjb')]

def add_locations_table(doc, form_data, visit_time, loc_key):
    p = doc.add_paragraph()
    para_spacing(p, 40, 40)
    add_run(p, '\u2022  Locations visited', bold=False, size_pt=11)

    t = doc.add_table(rows=5, cols=3)
    t.style = 'Table Grid'
    widths = [1604010, 1604010, 3209290]

    # Header
    hdr = t.rows[0]
    for i, (txt, w) in enumerate(zip(['Locations','Time','Remarks'], widths)):
        cell = hdr.cells[i]
        cell.text = txt
        set_col_width(cell, w)
        cell_margin(cell, 40, 40, 80, 80)
        for p2 in cell.paragraphs:
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(p2, 0, 0)
            for run in p2.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    row_height(hdr, 300)

    # Data rows
    for ri, (label, key) in enumerate(LOCS):
        row = t.rows[ri + 1]
        # Auto-fill time for matching hostel
        t_val = form_data.get(f'loc_{key}_time', '')
        if not t_val and key == loc_key:
            t_val = visit_time
        r_val = form_data.get(f'loc_{key}_remarks', '')
        vals = [label, t_val, r_val]
        for ci, (v, w) in enumerate(zip(vals, widths)):
            cell = row.cells[ci]
            cell.text = v
            set_col_width(cell, w)
            cell_margin(cell, 40, 40, 80, 80)
            for p2 in cell.paragraphs:
                para_spacing(p2, 0, 0)
                for run in p2.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
        row_height(row, 360)

    sp = doc.add_paragraph()
    para_spacing(sp, 60, 0)

# ── Status field (2-col table) ───────────────────────────────────────────────

def add_status_row(doc, label, value, col1_emu, col2_emu):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = label
    t.rows[0].cells[1].text = value or ''
    for i, (cell, w) in enumerate(zip(t.rows[0].cells, [col1_emu, col2_emu])):
        set_col_width(cell, w)
        cell_margin(cell, 40, 40, 80, 80)
        for p in cell.paragraphs:
            para_spacing(p, 0, 0)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    row_height(t.rows[0], 380)
    sp = doc.add_paragraph()
    para_spacing(sp, 30, 0)

# ── Suggestion box (multi-line bordered area) ────────────────────────────────

def add_suggestion_box(doc, label, value, rows=3):
    p = doc.add_paragraph()
    para_spacing(p, 60, 40)
    add_run(p, label, bold=True, size_pt=11)

    t = doc.add_table(rows=rows, cols=1)
    t.style = 'Table Grid'
    total_w = 6417310
    for ri in range(rows):
        cell = t.rows[ri].cells[0]
        set_col_width(cell, total_w)
        cell_margin(cell, 40, 40, 80, 80)
        if ri == 0:
            cell.text = value or ''
            for p2 in cell.paragraphs:
                para_spacing(p2, 0, 0)
                for run in p2.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
        row_height(t.rows[ri], 420)

    sp = doc.add_paragraph()
    para_spacing(sp, 40, 0)

# ── Signature area ───────────────────────────────────────────────────────────

def add_signature_antiragging(doc, faculty_name, phone):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    half = 3208655

    # Left: Signature of faculty
    cell_l = t.rows[0].cells[0]
    cell_l.text = 'Signature of the Faculty'
    set_col_width(cell_l, half)
    cell_margin(cell_l, 40, 300, 80, 80)
    for p in cell_l.paragraphs:
        para_spacing(p, 0, 0)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    row_height(t.rows[0], 1200)  # ~0.83" for signature space

    # Right: Cell number
    cell_r = t.rows[0].cells[1]
    set_col_width(cell_r, half)
    cell_margin(cell_r, 40, 300, 80, 80)
    p_r = cell_r.paragraphs[0]
    para_spacing(p_r, 0, 0)
    add_run(p_r, 'Cell No:-  ', bold=False, size_pt=11)
    add_run(p_r, phone or '', bold=False, size_pt=11)

    sp = doc.add_paragraph()
    para_spacing(sp, 80, 0)

def add_signature_mess(doc, faculty_name):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    half = 3208655

    for i, cell in enumerate(t.rows[0].cells):
        set_col_width(cell, half)
        cell_margin(cell, 40, 300, 80, 80)
        p = cell.paragraphs[0]
        para_spacing(p, 0, 0)
        add_run(p, 'Signature of Visiting faculty', bold=False, size_pt=11)
        p2 = cell.add_paragraph()
        para_spacing(p2, 120, 0)
        add_run(p2, f'{i+1})  {faculty_name if i == 0 else ""}', bold=False, size_pt=11)

    row_height(t.rows[0], 1200)

# ═══════════════════════════════════════════════════════════════════════════════
# ANTI-RAGGING FORM
# ═══════════════════════════════════════════════════════════════════════════════

def generate_anti_ragging(visit_data, form_data, output_path):
    doc = setup_page(Document())
    hostel_type = 'Girls' if visit_data.get('hostel_type') == 'girls' else 'Boys'
    year = visit_data.get('year', datetime.now().year)

    add_header(doc,
        committee_line='Anti-Ragging Committee',
        title_line=f'{hostel_type} Hostel Visit report {year}',
        hostel_type=hostel_type)

    add_date_name_rows(doc,
        date_str=fmt_date(visit_data.get('checkIn')),
        day_str=fmt_day(visit_data.get('checkIn')),
        time_str=fmt_time(visit_data.get('checkIn')),
        name_str=visit_data.get('faculty_name', ''),
        dept_str=visit_data.get('faculty_dept', ''))

    loc_key = hostel_to_loc_key(visit_data.get('hostel_name', ''))
    add_locations_table(doc, form_data, fmt_time(visit_data.get('checkIn')), loc_key)

    # Status fields
    add_status_row(doc, '\u2022  Status of Discipline',
        form_data.get('discipline_status',''), 1654810, 4762500)
    add_status_row(doc, '\u2022  Status of Cleanliness',
        form_data.get('cleanliness_status',''), 1654810, 4762500)
    add_status_row(doc, '\u2022  Overall Environment',
        form_data.get('environment_status',''), 1654810, 4762500)
    add_status_row(doc, '\u2022  Did you interact with senior students?',
        form_data.get('senior_interaction',''), 2734310, 3683000)
    add_status_row(doc, '\u2022  Did you interact with fresher student?',
        form_data.get('fresher_interaction',''), 2670810, 3746500)

    add_suggestion_box(doc, '\u2022  Suggestions related to Anti-ragging only',
        form_data.get('antiragging_suggestions',''), rows=3)
    add_suggestion_box(doc, '\u2022  Any other Suggestions',
        form_data.get('other_suggestions',''), rows=2)

    add_signature_antiragging(doc, visit_data.get('faculty_name',''), visit_data.get('faculty_phone',''))

    # Notes
    p_note = doc.add_paragraph()
    para_spacing(p_note, 60, 20)
    add_run(p_note, 'Note:', bold=True, size_pt=10)
    p1 = doc.add_paragraph()
    para_spacing(p1, 0, 0)
    add_run(p1, f'1) The faculty members should visit all the {hostel_type} Hostels (as applicable) on the premises', size_pt=10)
    p2 = doc.add_paragraph()
    para_spacing(p2, 0, 0)
    add_run(p2, '2) The visiting faculty should write feedback of issues related to only ragging in Hostel Visit Feedback Form and also submit the same to Anti Ragging Committee Coordinator by the next working day.', size_pt=10)

    # Helpline box
    sp = doc.add_paragraph()
    para_spacing(sp, 60, 0)
    t_h = doc.add_table(rows=1, cols=1)
    t_h.style = 'Table Grid'
    cell_h = t_h.rows[0].cells[0]
    cell_h.text = 'Anti-Ragging Helpline Number 1800-180-5522 (Toll-free) and email helpline@antiragging.net'
    set_col_width(cell_h, 6417310)
    cell_margin(cell_h, 60, 60, 80, 80)
    for p in cell_h.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, 0, 0)
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    doc.save(output_path)
    print(f"Saved: {output_path}", flush=True)

# ═══════════════════════════════════════════════════════════════════════════════
# MESS FEEDBACK FORM
# ═══════════════════════════════════════════════════════════════════════════════

def generate_mess_feedback(visit_data, form_data, output_path):
    doc = setup_page(Document())

    add_header(doc,
        committee_line='Hostel Mess Food Quality Inspection Committee',
        title_line='Feedback Form for Daily Inspection')

    add_date_name_rows(doc,
        date_str=fmt_date(visit_data.get('checkIn')),
        day_str=fmt_day(visit_data.get('checkIn')),
        time_str=fmt_time(visit_data.get('checkIn')),
        name_str=visit_data.get('faculty_name', ''),
        dept_str=visit_data.get('faculty_dept', ''))

    loc_key = hostel_to_loc_key(visit_data.get('hostel_name', ''))
    add_locations_table(doc, form_data, fmt_time(visit_data.get('checkIn')), loc_key)

    # Date + Meal row
    meal = form_data.get('meal_type', '')
    def meal_str(m):
        return f'{m} \u2713' if meal == m else m
    t_meal = doc.add_table(rows=1, cols=3)
    t_meal.style = 'Table Grid'
    t_meal.rows[0].cells[0].text = f'Date: {fmt_date(visit_data.get("checkIn",""))}'
    t_meal.rows[0].cells[1].text = ''
    t_meal.rows[0].cells[2].text = f'Meal: {meal_str("Breakfast")} / {meal_str("Lunch")} / {meal_str("Dinner")}'
    for i, (cell, w) in enumerate(zip(t_meal.rows[0].cells, [635000, 1587500, 4194810])):
        set_col_width(cell, w)
        cell_margin(cell, 40, 40, 80, 80)
        for p in cell.paragraphs:
            para_spacing(p, 0, 0)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    row_height(t_meal.rows[0], 360)
    sp = doc.add_paragraph(); para_spacing(sp, 40, 0)

    def yn(key):
        v = form_data.get(key, '')
        if v == 'Yes': return 'Yes \u2713 / No'
        if v == 'No':  return 'Yes / No \u2713'
        return 'Yes / No'

    def add_q(doc, text):
        p = doc.add_paragraph()
        para_spacing(p, 40, 0)
        add_run(p, text, size_pt=11)
        # Empty answer row table
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        cell = t.rows[0].cells[0]
        set_col_width(cell, 6417310)
        cell_margin(cell, 40, 40, 80, 80)
        cell.text = ''
        for p2 in cell.paragraphs:
            para_spacing(p2, 0, 0)
        row_height(t.rows[0], 380)
        sp = doc.add_paragraph(); para_spacing(sp, 30, 0)
        return t

    def add_q_inline(doc, text):
        p = doc.add_paragraph()
        para_spacing(p, 40, 0)
        add_run(p, text, size_pt=11)
        sp = doc.add_paragraph(); para_spacing(sp, 20, 0)

    add_q_inline(doc, f'1) Have you tasted food in the served meal? : {yn("tasted_food")}')

    p_q2 = doc.add_paragraph()
    para_spacing(p_q2, 40, 20)
    add_run(p_q2, '2) Menu items in the meal:', size_pt=11)
    t2 = doc.add_table(rows=1, cols=1)
    t2.style = 'Table Grid'
    cell2 = t2.rows[0].cells[0]
    set_col_width(cell2, 6417310)
    cell_margin(cell2, 40, 40, 80, 80)
    cell2.text = form_data.get('menu_items', '')
    for p2 in cell2.paragraphs:
        para_spacing(p2, 0, 0)
        for run in p2.runs: run.font.name='Times New Roman'; run.font.size=Pt(11)
    row_height(t2.rows[0], 420)
    sp = doc.add_paragraph(); para_spacing(sp, 30, 0)

    add_q_inline(doc, f'3) Has the cleanliness of the dining hall been maintained? : {yn("cleanliness")}')
    add_q_inline(doc, f'4) Were the empty Plates/Spoons neatly cleaned? : {yn("plates_clean")}')
    add_q_inline(doc, f'5) Was the food served hot? : {yn("food_hot")}')

    p_q6 = doc.add_paragraph()
    para_spacing(p_q6, 40, 20)
    add_run(p_q6, '6) Write your detailed remarks about the taste and condition of the food (meal) you have tasted :', size_pt=11)
    t6 = doc.add_table(rows=2, cols=1)
    t6.style = 'Table Grid'
    for ri in range(2):
        cell6 = t6.rows[ri].cells[0]
        set_col_width(cell6, 6417310)
        cell_margin(cell6, 40, 40, 80, 80)
        if ri == 0:
            cell6.text = form_data.get('food_remarks', '')
            for p2 in cell6.paragraphs:
                para_spacing(p2, 0, 0)
                for run in p2.runs: run.font.name='Times New Roman'; run.font.size=Pt(11)
        row_height(t6.rows[ri], 400)
    sp = doc.add_paragraph(); para_spacing(sp, 30, 0)

    ovf = form_data.get('overall_feedback', '')
    def ovf_str(v):
        return f'{v} \u2713' if ovf == v else v
    add_q_inline(doc, f'7) Overall Feedback about the food: {ovf_str("Satisfactory")} / {ovf_str("Needs improvement")}')

    p_q8 = doc.add_paragraph()
    para_spacing(p_q8, 40, 20)
    add_run(p_q8, '8) Please suggest areas of improvement:', size_pt=11)
    t8 = doc.add_table(rows=2, cols=1)
    t8.style = 'Table Grid'
    for ri in range(2):
        cell8 = t8.rows[ri].cells[0]
        set_col_width(cell8, 6417310)
        cell_margin(cell8, 40, 40, 80, 80)
        if ri == 0:
            cell8.text = form_data.get('improvement_suggestions', '')
            for p2 in cell8.paragraphs:
                para_spacing(p2, 0, 0)
                for run in p2.runs: run.font.name='Times New Roman'; run.font.size=Pt(11)
        row_height(t8.rows[ri], 400)
    sp = doc.add_paragraph(); para_spacing(sp, 60, 0)

    add_signature_mess(doc, visit_data.get('faculty_name', ''))
    doc.save(output_path)
    print(f"Saved: {output_path}", flush=True)

# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python3 generate_form.py <form_type> <json_data> <output_path>")
        sys.exit(1)

    form_type   = sys.argv[1]
    data        = json.loads(sys.argv[2])
    output_path = sys.argv[3]

    visit = data.get('visit', {})
    form  = data.get('formData', {})

    if form_type == 'anti_ragging':
        generate_anti_ragging(visit, form, output_path)
    elif form_type == 'mess_feedback':
        generate_mess_feedback(visit, form, output_path)
    else:
        print(f"Unknown form type: {form_type}")
        sys.exit(1)